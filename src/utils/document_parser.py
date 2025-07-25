import logging
import io
import os
import shutil
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import PyPDF2
import pytesseract
from PIL import Image
import docx
import re
from datetime import datetime

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser:
    """
    Parses various medical document formats including PDFs, images, and text files.
    Extracts text and structured information from medical documents.
    """
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._parse_pdf,
            '.png': self._parse_image,
            '.jpg': self._parse_image,
            '.jpeg': self._parse_image,
            '.tiff': self._parse_image,
            '.bmp': self._parse_image,
            '.docx': self._parse_docx,
            '.txt': self._parse_text
        }
        
        # Check if Tesseract is available
        self.tesseract_available = self._check_tesseract()
        
        # Common medical document patterns
        self.patterns = {
            'date': r'\b(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})\b',
            'patient_name': r'Patient(?:\s+Name)?:\s*([A-Za-z\s]+)',
            'doctor_name': r'(?:Dr\.|Doctor|Physician):\s*([A-Za-z\s]+)',
            'medical_record_number': r'MRN?:\s*(\d+)',
            'lab_result': r'([A-Za-z\s]+):\s*(\d+\.?\d*)\s*([A-Za-z/%]+)?'
        }
    
    def _check_tesseract(self) -> bool:
        """Check if Tesseract OCR is installed and available."""
        try:
            # Check if tesseract command exists
            if shutil.which('tesseract') is None:
                logger.warning("Tesseract OCR not found in PATH")
                return False
            
            # Try to get tesseract version
            pytesseract.get_tesseract_version()
            logger.info("Tesseract OCR is available")
            return True
        except Exception as e:
            logger.warning(f"Tesseract OCR not available: {str(e)}")
            return False
    
    async def parse_document(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parses a medical document and extracts text and structured information.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}. Supported formats: {', '.join(self.supported_formats.keys())}")
        
        logger.info(f"Parsing document: {file_path}")
        
        # Check if OCR is needed and available
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'] and not self.tesseract_available:
            raise RuntimeError(
                "Tesseract OCR is required to read image files but is not installed. "
                "Please install it with: brew install tesseract (macOS) or "
                "sudo apt-get install tesseract-ocr (Linux)"
            )
        
        # Parse the document
        parser_func = self.supported_formats[file_extension]
        raw_text = await parser_func(file_path)
        
        # Check if parsing returned an error message
        if raw_text.startswith("Unable to") and not self.tesseract_available:
            raise RuntimeError(
                "Document parsing failed. Tesseract OCR is required for this document type but is not installed. "
                "Please install it with: brew install tesseract (macOS) or "
                "sudo apt-get install tesseract-ocr (Linux)"
            )
        
        # Extract structured information
        metadata = self._extract_metadata(raw_text)
        
        # Clean and structure the text
        cleaned_text = self._clean_text(raw_text)
        
        return {
            "file_path": str(file_path),
            "file_type": file_extension,
            "raw_text": raw_text,
            "cleaned_text": cleaned_text,
            "metadata": metadata,
            "extracted_values": self._extract_medical_values(cleaned_text),
            "sections": self._identify_sections(cleaned_text),
            "parsing_notes": self._get_parsing_notes(file_extension, raw_text)
        }
    
    def _get_parsing_notes(self, file_extension: str, raw_text: str) -> List[str]:
        """Generate notes about the parsing process."""
        notes = []
        
        if file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
            if self.tesseract_available:
                notes.append("Used OCR to extract text from image")
            else:
                notes.append("OCR unavailable - image parsing may be limited")
        
        if file_extension == '.pdf' and len(raw_text) < 100:
            notes.append("Short text extracted - document may be image-based or encrypted")
            if not self.tesseract_available:
                notes.append("Install Tesseract OCR for better PDF image processing")
        
        if raw_text.startswith("Unable to"):
            notes.append("Text extraction failed - check document format and system dependencies")
        
        return notes
    
    async def _parse_pdf(self, file_path: Path) -> str:
        """Extracts text from PDF files."""
        text_content = []
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    text_content.append(f"--- Page {page_num + 1} ---\n{text}")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            # Try OCR as fallback
            return await self._parse_pdf_with_ocr(file_path)
    
    async def _parse_pdf_with_ocr(self, file_path: Path) -> str:
        """Fallback OCR for PDFs that can't be parsed normally."""
        if not self.tesseract_available:
            return ("Unable to extract text from PDF: This appears to be a scanned/image-based PDF. "
                   "Tesseract OCR is required but not installed. Please install it with: "
                   "brew install tesseract (macOS) or sudo apt-get install tesseract-ocr (Linux)")
        
        try:
            # Convert PDF pages to images and OCR them
            from pdf2image import convert_from_path
            
            images = convert_from_path(file_path)
            text_content = []
            
            for i, image in enumerate(images):
                text = pytesseract.image_to_string(image, config='--psm 3')
                text_content.append(f"--- Page {i + 1} ---\n{text}")
            
            extracted_text = "\n\n".join(text_content)
            
            # Check if OCR actually extracted meaningful text
            if len(extracted_text.strip()) < 10:
                return ("Unable to extract meaningful text from PDF: "
                       "The document may be low quality, encrypted, or contain non-text content.")
            
            return extracted_text
            
        except Exception as e:
            logger.error(f"Error in PDF OCR: {str(e)}")
            return f"Unable to extract text from PDF: OCR failed with error: {str(e)}"
    
    async def _parse_image(self, file_path: Path) -> str:
        """Extracts text from image files using OCR."""
        if not self.tesseract_available:
            return ("Unable to extract text from image: Tesseract OCR is required but not installed. "
                   "Please install it with: brew install tesseract (macOS) or "
                   "sudo apt-get install tesseract-ocr (Linux)")
        
        try:
            image = Image.open(file_path)
            
            # Preprocess image for better OCR
            # Convert to grayscale
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance image for better OCR (optional improvements)
            # Could add more preprocessing here if needed
            
            # Extract text using OCR with optimized settings
            text = pytesseract.image_to_string(image, config='--psm 3 -c tessedit_char_whitelist=0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz .,:-/()%')
            
            # Check if OCR extracted meaningful text
            if len(text.strip()) < 5:
                return ("Unable to extract meaningful text from image: "
                       "The image may be low quality, blurry, or contain non-text content. "
                       "Try using a higher quality scan or image.")
            
            return text
            
        except Exception as e:
            logger.error(f"Error parsing image: {str(e)}")
            return f"Unable to extract text from image: {str(e)}"
    
    async def _parse_docx(self, file_path: Path) -> str:
        """Extracts text from Word documents."""
        try:
            doc = docx.Document(file_path)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            return "\n".join(text_content)
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {str(e)}")
            return "Unable to extract text from document"
    
    async def _parse_text(self, file_path: Path) -> str:
        """Reads plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {str(e)}")
            return "Unable to read text file"
    
    def _clean_text(self, text: str) -> str:
        """Cleans and normalizes extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but keep medical symbols
        text = re.sub(r'[^\w\s\-.,/:;()%°]', ' ', text)
        
        # Fix common OCR errors
        ocr_corrections = {
            r'\bl\b': '1',  # lowercase L to 1
            r'\bO\b': '0',  # uppercase O to 0
            r'l\/': '1/',    # l/ to 1/
        }
        
        for pattern, replacement in ocr_corrections.items():
            text = re.sub(pattern, replacement, text)
        
        return text.strip()
    
    def _extract_metadata(self, text: str) -> Dict[str, Any]:
        """Extracts metadata from document text."""
        metadata = {}
        
        # Extract dates
        date_matches = re.findall(self.patterns['date'], text)
        if date_matches:
            metadata['dates'] = date_matches
            # Try to identify report date
            report_date_pattern = r'(?:Report|Test|Lab)\s+Date:\s*(\d{1,2}[-/]\d{1,2}[-/]\d{2,4})'
            report_date = re.search(report_date_pattern, text, re.IGNORECASE)
            if report_date:
                metadata['report_date'] = report_date.group(1)
        
        # Extract patient name
        patient_match = re.search(self.patterns['patient_name'], text, re.IGNORECASE)
        if patient_match:
            metadata['patient_name'] = patient_match.group(1).strip()
        
        # Extract doctor name
        doctor_match = re.search(self.patterns['doctor_name'], text, re.IGNORECASE)
        if doctor_match:
            metadata['doctor_name'] = doctor_match.group(1).strip()
        
        # Extract MRN
        mrn_match = re.search(self.patterns['medical_record_number'], text, re.IGNORECASE)
        if mrn_match:
            metadata['mrn'] = mrn_match.group(1)
        
        return metadata
    
    def _extract_medical_values(self, text: str) -> List[Dict[str, Any]]:
        """Extracts medical test values from text."""
        values = []
        
        # Common medical test patterns - ordered from most specific to least specific
        test_patterns = [
            # Cholesterol tests - specific ones first to avoid conflicts
            (r'LDL\s*(?:Cholesterol)?[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'LDL Cholesterol'),
            (r'HDL\s*(?:Cholesterol)?[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'HDL Cholesterol'),
            (r'(?:Total\s+)?Cholesterol[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'Total Cholesterol'),
            (r'Triglycerides[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'Triglycerides'),
            
            # Blood tests
            (r'(?:Hemoglobin|Hgb|HGB)[:\s]+(\d+\.?\d*)\s*(g/dL|g/dl)?', 'Hemoglobin'),
            (r'(?:Glucose|GLU|Blood\s+Glucose)[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'Glucose'),
            (r'(?:Blood\s+)?Pressure[:\s]+(\d+/\d+)\s*(mmHg)?', 'Blood Pressure'),
            
            # Complete Blood Count
            (r'(?:WBC|White\s+Blood\s+Cells?)[:\s]+(\d+\.?\d*)\s*(K/uL|cells/uL|K/μL)?', 'White Blood Cells'),
            (r'(?:RBC|Red\s+Blood\s+Cells?)[:\s]+(\d+\.?\d*)\s*(M/uL|M/μL)?', 'Red Blood Cells'),
            (r'Platelets?[:\s]+(\d+\.?\d*)\s*(K/uL|K/μL)?', 'Platelets'),
            (r'Hematocrit[:\s]+(\d+\.?\d*)\s*%?', 'Hematocrit'),
            
            # Metabolic panel
            (r'Creatinine[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'Creatinine'),
            (r'BUN[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'BUN'),
            (r'(?:Sodium|Na)[:\s]+(\d+\.?\d*)\s*(mEq/L|mmol/L)?', 'Sodium'),
            (r'(?:Potassium|K)[:\s]+(\d+\.?\d*)\s*(mEq/L|mmol/L)?', 'Potassium'),
            (r'(?:Chloride|Cl)[:\s]+(\d+\.?\d*)\s*(mEq/L|mmol/L)?', 'Chloride'),
            
            # Liver function
            (r'(?:ALT|SGPT)[:\s]+(\d+\.?\d*)\s*(U/L|IU/L)?', 'ALT'),
            (r'(?:AST|SGOT)[:\s]+(\d+\.?\d*)\s*(U/L|IU/L)?', 'AST'),
            (r'(?:Total\s+)?Bilirubin[:\s]+(\d+\.?\d*)\s*(mg/dL|mg/dl)?', 'Bilirubin'),
            (r'(?:Alkaline\s+)?Phosphatase[:\s]+(\d+\.?\d*)\s*(U/L|IU/L)?', 'Alkaline Phosphatase'),
            
            # Thyroid function
            (r'TSH[:\s]+(\d+\.?\d*)\s*(mIU/L|μIU/ml)?', 'TSH'),
            (r'T4[:\s]+(\d+\.?\d*)\s*(μg/dL|pmol/L)?', 'T4'),
            (r'T3[:\s]+(\d+\.?\d*)\s*(ng/dL|pmol/L)?', 'T3'),
            
            # Diabetes markers
            (r'(?:HbA1c|A1C|Hemoglobin\s+A1c)[:\s]+(\d+\.?\d*)\s*%?', 'HbA1c'),
            
            # Inflammation markers
            (r'(?:ESR|Sed\s+Rate)[:\s]+(\d+\.?\d*)\s*(mm/hr)?', 'ESR'),
            (r'(?:CRP|C-Reactive\s+Protein)[:\s]+(\d+\.?\d*)\s*(mg/L|mg/dL)?', 'CRP'),
            
            # Vitamins
            (r'Vitamin\s+D[:\s]+(\d+\.?\d*)\s*(ng/mL|nmol/L)?', 'Vitamin D'),
            (r'Vitamin\s+B12[:\s]+(\d+\.?\d*)\s*(pg/mL|pmol/L)?', 'Vitamin B12'),
            (r'Folate[:\s]+(\d+\.?\d*)\s*(ng/mL|nmol/L)?', 'Folate'),
        ]
        
        for pattern, test_name in test_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                value_dict = {
                    'test_name': test_name,
                    'value': match.group(1),
                    'unit': match.group(2) if match.lastindex >= 2 else None,
                    'position': match.start()
                }
                values.append(value_dict)
        
        # Sort by position in document
        values.sort(key=lambda x: x['position'])
        
        # Remove position from final output
        for v in values:
            v.pop('position', None)
        
        return values
    
    def _identify_sections(self, text: str) -> Dict[str, str]:
        """Identifies common sections in medical documents."""
        sections = {}
        
        # Common section headers
        section_patterns = {
            'chief_complaint': r'(?:Chief\s+Complaint|CC|Reason\s+for\s+Visit)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'history': r'(?:History|HPI|Past\s+Medical\s+History)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'medications': r'(?:Medications?|Current\s+Medications?|Meds)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'allergies': r'(?:Allergies|Drug\s+Allergies)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'assessment': r'(?:Assessment|Impression|Diagnosis)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'plan': r'(?:Plan|Treatment\s+Plan|Recommendations?)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'lab_results': r'(?:Lab\s+Results?|Laboratory\s+Results?)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)',
            'vital_signs': r'(?:Vital\s+Signs?|Vitals)[:\s]+(.*?)(?=\n[A-Z]|\n\n|\Z)'
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                content = match.group(1).strip()
                # Limit section length
                if len(content) > 1000:
                    content = content[:997] + "..."
                sections[section_name] = content
        
        return sections
    
    async def extract_medications_from_document(self, file_path: Union[str, Path]) -> List[Dict[str, str]]:
        """
        Specifically extracts medication information from a document.
        
        Returns:
            List of medications with dosage and frequency
        """
        document_data = await self.parse_document(file_path)
        text = document_data['cleaned_text']
        
        medications = []
        
        # Medication patterns
        med_patterns = [
            # Standard format: Drug name dose frequency
            r'([A-Za-z]+(?:\s+[A-Za-z]+)?)\s+(\d+\.?\d*\s*(?:mg|mcg|g|ml|units?))\s+([A-Za-z\s]+(?:daily|BID|TID|QID|PRN))',
            # Alternative format
            r'([A-Za-z]+(?:\s+[A-Za-z]+)?)\s+-\s+(\d+\.?\d*\s*(?:mg|mcg|g|ml|units?))',
        ]
        
        for pattern in med_patterns:
            matches = re.finditer(pattern, text, re.IGNORECASE)
            for match in matches:
                med = {
                    'name': match.group(1).strip(),
                    'dosage': match.group(2).strip() if match.lastindex >= 2 else 'Unknown',
                    'frequency': match.group(3).strip() if match.lastindex >= 3 else 'As directed'
                }
                
                # Avoid duplicates
                if not any(m['name'].lower() == med['name'].lower() for m in medications):
                    medications.append(med)
        
        return medications